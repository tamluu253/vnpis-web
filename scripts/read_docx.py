import sys
import io
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')
from docx import Document

def read_docx(path):
    try:
        doc = Document(path)
        print("Text:")
        for p in doc.paragraphs:
            if p.text.strip():
                print(p.text.strip())
        print("Tables:")
        for table in doc.tables:
            for row in table.rows:
                row_data = [cell.text.strip() for cell in row.cells]
                print(row_data)
        
        # Check for images
        images = []
        for rel in doc.part.rels.values():
            if "image" in rel.target_ref:
                images.append(rel.target_ref)
        print(f"Contains {len(images)} images.")
    except Exception as e:
        print(f"Error: {e}")

if __name__ == "__main__":
    read_docx(sys.argv[1])
