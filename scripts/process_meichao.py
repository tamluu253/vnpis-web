import os
import io
import sys
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')
import json
import uuid
import shutil
import docx
import zipfile

def extract_image(docx_path, out_dir, prefix):
    # Extract the first image from docx
    try:
        with zipfile.ZipFile(docx_path, 'r') as z:
            image_names = [n for n in z.namelist() if n.startswith('word/media/') and '.' in n]
            if image_names:
                img_name = image_names[0]
                ext = img_name.split('.')[-1]
                out_name = f"{prefix}.{ext}"
                out_path = os.path.join(out_dir, out_name)
                with open(out_path, 'wb') as f:
                    f.write(z.read(img_name))
                return f"/images/printers/mc/{out_name}"
    except Exception as e:
        print(f"Failed to extract image from {docx_path}: {e}")
    return ""

def process_directory(directory_path, output_json, output_image_dir):
    os.makedirs(output_image_dir, exist_ok=True)
    printers = []
    
    for filename in os.listdir(directory_path):
        if not filename.endswith('.docx') or filename.startswith('~'):
            continue
            
        filepath = os.path.join(directory_path, filename)
        print(f"Processing {filename}...")
        try:
            doc = docx.Document(filepath)
            text_lines = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
            
            title = text_lines[0] if text_lines else filename.replace('.docx', '')
            specs = []
            
            for line in text_lines[1:]:
                if ':' in line or '：' in line:
                    specs.append(line)
            
            # Simple heuristic
            if len(specs) == 0:
                for table in doc.tables:
                    for row in table.rows:
                        row_data = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                        if row_data:
                            specs.append(" : ".join(row_data))
                            
            printer_id = f"mc-{uuid.uuid4().hex[:6]}"
            img_url = extract_image(filepath, output_image_dir, printer_id)
            
            printers.append({
                "id": printer_id,
                "name": title,
                "desc": "Máy in Meichao (MC series) - " + filename,
                "features": specs[:5], # Take top 5 specs as features
                "specs": specs,
                "image": img_url
            })
        except Exception as e:
            print(f"Error reading {filename}: {e}")
            
    with open(output_json, 'w', encoding='utf-8') as f:
        json.dump(printers, f, indent=2, ensure_ascii=False)
    print(f"Saved {len(printers)} printers to {output_json}")

if __name__ == "__main__":
    src_dir = r"C:\Users\TL\Documents\Disk D\cosota\## Ali - Cons supplier\Meichao\catalog"
    out_json = r"C:\Users\TL\Documents\vnpis-web\src\data\mc-printers.json"
    out_img = r"C:\Users\TL\Documents\vnpis-web\public\images\printers\mc"
    process_directory(src_dir, out_json, out_img)
